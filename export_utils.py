from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document

def export_pdf(content, filename):

    doc = SimpleDocTemplate(filename)

    styles = getSampleStyleSheet()

    story = [
        Paragraph(
            content.replace("\n", "<br/>"),
            styles["BodyText"]
        )
    ]

    doc.build(story)

def export_docx(content, filename):

    doc = Document()

    doc.add_heading(
        "Generated Course",
        level=1
    )

    doc.add_paragraph(content)

    doc.save(filename)